import string
from docx.api import Document
from lxml import etree

def getRequest(arr):
	string = ''
	for word in arr:
		string += word + '____'
	return string.split('____')	

def readDocxParagraph(name):
	str = ''
	document = Document(name)
	for para in document.paragraphs:
		str = str + para.text + u" endpara "
		
	str = str.lower()
	return str
	
def readDocxTable(name):
	str = ''
	document = Document(name)
	tables = document.tables
	for table in tables:
		for row in table.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					str = str + paragraph.text + u" endpara "
	
	str = str.lower()
	return str
	
def readTxt(path):
	return open("C:\\xampp\\htdocs\\Plagiarism\\python\\keywords\\" + path + ".txt", "r")
	
def getKeyword(list, s):
	result = ''
	for line in list:
		key = line.lstrip().rstrip().decode('utf_8')
		keyReplace = key.replace(' ', '_')
		sum = s.count(key, 0 , len(s)) # tong so lan xuat hien
		if(sum > 0):
			keyReplace = keyReplace.replace(keyReplace, "(" + keyReplace + "|" + str(sum) +")")
			result = result + keyReplace + ' '
	return result.lstrip().rstrip()

def getContent(list, s):
	for line in list:
		key = line.lstrip().rstrip().decode('utf_8')
		s = s.replace(key, key.replace(' ', '_'))
	return s.lstrip().rstrip()
 
def convertXML(list, key):
	context = etree.Element("context")
	contentPara = etree.SubElement(context, "content-param")
	
	for s in list:
		para = etree.SubElement(contentPara, "para")
		keywords = etree.SubElement(para, "keywords")
		content = etree.SubElement(para, 'content')
		keywords.text = getKeyword(readTxt(key), s)
		content.text = getContent(readTxt(key), s)
		
	return etree.tostring(context, pretty_print=True)
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	